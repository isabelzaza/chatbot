import requests
import json
import streamlit as st
import PyPDF2
import docx
import io
import gspread

# Define the inventory questions
INVENTORY_QUESTIONS = {
    "Q1": {"question": "Instructor Name", "format": "text"},
    "Q2": {"question": "Course Number", "format": "text"},
    "Q3": {"question": "Semester and Year", "format": "text"},
    "Q4": {"question": "Number of Students", "format": "number"},
    "Q5": {"question": "Do you give students a list of the topics that will be covered in the course?", "format": "y/n"},
    "Q6": {"question": "Do you provide a list of general skills students should develop (like critical thinking)?", "format": "y/n"},
    "Q7": {"question": "Do you explain what specific skills students should learn from specific topics (e.g., from a given activity or lecture)?", "format": "y/n"},
    "Q8": {"question": "Do you articulate a policy regarding permissible use of AI tools/LLMs for this class, beyond saying all use of AI tools in assignments is banned?", "format": "y/n"},
    "Q9": {"question": "Do you use a platform for class discussions online?", "format": "y/n"},
    "Q10": {"question": "Do you use a course website like Brightspace to share materials?", "format": "y/n"},
    "Q11": {"question": "Do you provide solutions to homework assignments?", "format": "y/n"},
    "Q12": {"question": "Do you share examples of how to solve problems (e.g., step-by-step guides)?", "format": "y/n"},
    "Q13": {"question": "Do you give practice tests or past exam papers?", "format": "y/n"},
    "Q14": {"question": "Do you share your lecture slides, notes or recording of lectures with students (either before or after class)?", "format": "y/n"},
    "Q15": {"question": "Do you provide other materials, like reading or study aids?", "format": "y/n"},
    "Q16": {"question": "Do you give students access to articles or research related to the course?", "format": "y/n"},
    "Q17": {"question": "Do you share examples of excellent student papers or projects?", "format": "y/n"},
    "Q18": {"question": "Do you provide grading guides/rubrics to explain how assignments will be marked?", "format": "y/n"},
    "Q19": {"question": "Do you regularly use a strategy to elicit student questions in class -beyond saying are there any questions and moving on?", "format": "y/n"},
    "Q20": {"question": "In a typical class (excluding special days like exams or student presentations), what is the proportion of time for which students work in small groups? e.g., 50 would mean 50%", "format": "percentage (0 to 100)"},
    "Q21": {"question": "In a typical class (excluding special days), what is the proportion of time for which you lecture? e.g., 50 would mean 50%", "format": "percentage (0 to 100)"},
    "Q22": {"question": "What is the maximum time you typically spend lecturing continuously before switching to a completely different activity? (in minutes)", "format": "number (minutes)"},
    "Q23": {"question": "How often do you use demonstrations, simulations, or videos?", 
            "format": "choice: never/rarely/once in a while/every week/every class"},
    "Q24": {"question": "If you show a video, do you provide students with detailed instructions of what to look for in the video, and follow-up with a discussion?", "format": "y/n"},
    "Q25": {"question": "How often do you talk about why the material might be useful or interesting from a student's point of view?", 
            "format": "choice: never/rarely/once in a while/every week/every class"},
    "Q26": {"question": "Do you use a response system (e.g., Top hat) for ungraded activities?", 
            "format": "choice: never/rarely/once in a while/every week/every class"},
    "Q27": {"question": "Do you use a response system (e.g., Top hat) for graded activities?", 
            "format": "choice: never/rarely/once in a while/every week/every class"},
    "Q28": {"question": "Do you give homework or practice problems that do not count towards grade?", "format": "y/n"},
    "Q29": {"question": "Do you give homework or problems that count towards grade?", "format": "y/n"},
    "Q30": {"question": "Do you assign a project or paper where students have some choice about the topic?", "format": "y/n"},
    "Q31": {"question": "Do you encourage students to work together on individual assignments?", "format": "y/n"},
    "Q32": {"question": "Do you give group assignments?", "format": "y/n"},
    "Q33": {"question": "Do you ask students for anonymous feedback outside of end of term course reviews", "format": "y/n"},
    "Q34": {"question": "Do you allow students to revise their work based on feedback?", "format": "y/n"},
    "Q35": {"question": "Do students have access to their graded exams and other assignments?", "format": "y/n"},
    "Q36": {"question": "Do you share answer keys for graded exams and other assignments?", "format": "y/n"},
    "Q37": {"question": "Do you require students to include a statement regarding their use of AI tools/LLMs on submitted assignments, or submit some kind of record to delineate which parts of the assignment represent their own intellectual contributions versus those of generative AI?", "format": "y/n"},
    "Q38": {"question": "How often do you have students consult with you outside of class?", "format": "choice: never/rarely/once a week/several times a week"},
    "Q39": {"question": "Do you give a test at the beginning of the course to see what students already know?", "format": "y/n"},
    "Q40": {"question": "Do you use a pre-and-post test to measure how much students learn in the course?", "format": "y/n"},
    "Q41": {"question": "Do you ask students about their interest or feelings about the subject before and after the course?", "format": "y/n"},
    "Q42": {"question": "Do you give students some control over their learning, like choosing topics or how they will be graded?", "format": "y/n"},
    "Q43": {"question": "Do you try new teaching methods or materials and measure how well they work?", 
            "format": "choice: never/rarely/sometimes/often"},
    "Q44": {"question": "Do you have graduate TAs or undergraduate LAs for this course?", "format": "y/n"},
    "Q45": {"question": "Do you provide TAs/LAs with some training on teaching methods?", 
            "format": "choice: not applicable/no/yes"},
    "Q46": {"question": "Do you meet with TAs/LAs regularly to talk about teaching and how students are doing?", 
            "format": "choice: not applicable/no/yes"},
    "Q47": {"question": "Do you use teaching materials from other instructors?", "format": "y/n"},
    "Q48": {"question": "Do you use some of the same teaching materials as other instructors of the same course in your department?", "format": "y/n"},
    "Q49": {"question": "Do you talk with colleagues about how to teach this course?", "format": "y/n"},
    "Q50": {"question": "Relevant to this course, did you read articles or attend workshops to improve your teaching?", "format": "y/n"},
    "Q51": {"question": "Have you observed another instructor's class to get ideas?", "format": "y/n"},
    "Q52": {"question": "Do you have any comments about this inventory?",
            "format": "text"}
}

# Section Definitions - Balanced sections with 6-7 questions each
SECTIONS = {
    "Section 1": ["Q1", "Q2", "Q3", "Q4", "Q6", "Q7"],
    "Section 2": ["Q8", "Q9", "Q11", "Q12", "Q13", "Q14"],
    "Section 3": ["Q15", "Q16", "Q17", "Q18", "Q19", "Q20"],
    "Section 4": ["Q21", "Q22", "Q23", "Q24", "Q25", "Q26"],
    "Section 5": ["Q27", "Q28", "Q29", "Q30", "Q31", "Q32"],
    "Section 6": ["Q33", "Q34", "Q35", "Q36", "Q37", "Q38"],
    "Section 7": ["Q39", "Q40", "Q41", "Q42", "Q43", "Q45"],
    "Section 8": ["Q46", "Q47", "Q48", "Q49", "Q50", "Q51", "Q52"]
}
# LLM Prompt Template
INVENTORY_PROMPT = """
Analyze the provided document(s) to answer ALL questions from the Vanderbilt Psychology Teaching Inventory (Q1-Q52). 
Consider ALL provided documents together.

FORMAT GUIDELINES:
For each answer you find, provide:
1. Question number (Q1-Q51)
2. Question text
3. Answer - ONLY if explicitly found in documents
4. Evidence: Quote "exactly as appears" (From Document 1/2)

CORE QUESTIONS:

Q1 (Instructor Name):
- Look for "Instructor" or containing "Teacher"/"Professor" followed by a name

Q2 (Course Number):
- Format: "PSY ####" (or NSC/BSCI etc.)
- Appears near top of syllabus
- Do not use course title or other text

Q3 (Semester/Year):
- Look for full semester name + four-digit year
- Could say "Spring", "Fall" or "Summer" followed by a year
- Check document names (e.g., "S2021", "F21")
- Must have both semester and year explicitly stated

Q6 (General Skills):
Look for broad skills in:
- Course objectives/goals
- Learning outcomes
Examples:
- Critical thinking
- Analysis
- Problem-solving
- Research methods

Q7 (Topic-Specific Skills):
Look for skills explicitly tied to:
- Specific topics
- Specific assignments
- Course modules
Must connect skills to particular content/activities

Q8 (AI Policy):
Look for
- "Artificial Intelligence" or "AI" and "policy"
- "ChatGPT"
- "Generative AI"
- "LLM" or "Large language models"

Q9 (Online Discussions):
Look for:
- Discussion board (likely on Brightspace)
- posting or sharing comments
- Online forums
- Required posting/responses
- References to online participation
- Mentions of posting/replying in participation sections

Q13 (practice tests)
Look for:
-	“practice test”, “old test”

Q14 (sharing material):
Look for information on slides or lecture notes that
- "will be posted"
- "will be made available"

Q15 (study aids):
Look for
-	“reading guide”, “study guide”

Q16 (articles)
Look for
-	“primary articles”, “research articles”

Q17 (student work examples)
Look for
-	“examples of good work”, “example papers”, “examples from prior semesters”

Q18 (grading rubrics)"
Look for mention of "rubric"

Q26/27 (Response Systems):
Look for:
- TopHat, iClicker, etc.
- Q26: Ungraded use ("practice", "demos")
- Q27: Graded use ("points", "required")

Q28/29 (Homework):
- Q28: Ungraded practice/optional work
- Q29: Graded assignments/required work
Check if points/grades assigned

Q30 (Project with options):
Look for:
-discussion of a project or assignment, with "option(s)" or "choice(s)"

Q32 (Team work):
Look for:
-Group work; group assignment

Q33 (anonymous survey):
Look for:
- "anonymous feedback"
- "mid-term feedback" or "mid-term survey"

Q34 (revision based on feedback):
Look for:
- mention of "revision" or "revise"
- feedback

Q45/Q46 (Teaching Assistants):
Look for:
- "Teaching assistants" sections
- Lists of TA names/emails
- Mentions of "TA" or "LA"
- If no TAs/LAs are mentioned, then both Q45 and Q46 should be "not applicable"

IMPORTANT: Only provide answers with explicit evidence from the documents. Do not infer or guess.

The documents:
{documents}

"""

# File Processing Functions
def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def process_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return None, None
    
    try:
        file_bytes = uploaded_file.getvalue()
        filename = uploaded_file.name  # Get the filename
        
        if uploaded_file.type == "application/pdf":
            text = read_pdf(io.BytesIO(file_bytes))
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = read_docx(io.BytesIO(file_bytes))
        else:
            st.error("Unsupported file type. Please upload a PDF or Word document.")
            return None, None
            
        return text, filename
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None, None

def save_to_google_sheets(answers):
    """Save answers to Google Sheets"""
    try:
        # Get credentials from Streamlit secrets
        credentials = {
            "type": st.secrets["gcp_service_account"]["type"],
            "project_id": st.secrets["gcp_service_account"]["project_id"],
            "private_key_id": st.secrets["gcp_service_account"]["private_key_id"],
            "private_key": st.secrets["gcp_service_account"]["private_key"],
            "client_email": st.secrets["gcp_service_account"]["client_email"],
            "client_id": st.secrets["gcp_service_account"]["client_id"],
            "auth_uri": st.secrets["gcp_service_account"]["auth_uri"],
            "token_uri": st.secrets["gcp_service_account"]["token_uri"],
            "auth_provider_x509_cert_url": st.secrets["gcp_service_account"]["auth_provider_x509_cert_url"],
            "client_x509_cert_url": st.secrets["gcp_service_account"]["client_x509_cert_url"],
            "universe_domain": st.secrets["gcp_service_account"]["universe_domain"]
        }
        
        # Create row of answers in correct order
        row = []
        for i in range(1, 53):  # Updated to include Q52
            q_id = f"Q{i}"
            # Mark Q5, Q10, and Q44 as "notasked"
            if q_id in ["Q5", "Q10", "Q44"]:
                row.append("notasked")
            else:
                row.append(answers.get(q_id, ""))
            
        # Connect to Google Sheets
        gc = gspread.service_account_from_dict(credentials)
        sheet_url = st.secrets["google_sheets"]["sheet_url"]
        spreadsheet = gc.open_by_url(sheet_url)
        worksheet = spreadsheet.sheet1  # Gets the first sheet
        
        # Append row to sheet
        worksheet.append_row(row)
        return True

    except Exception as e:
        st.error(f"Error saving to Google Sheets: {str(e)}")
        return False

# LLM Response Parsing
def parse_llm_response(response_text):
    """
    Parse the LLM's response into a dictionary of answers and evidence
    """
    answers = {}
    evidence = {}
    current_question = None
    current_evidence = None

    # DEBUG: Show what we're parsing
    st.write("🔍 **DEBUG: Parsing LLM Response**")
    st.write(f"- Total response length: {len(response_text)} characters")
    st.write(f"- Total lines: {len(response_text.split(chr(10)))}")

    try:
        lines = response_text.split('\n')

        # Track parsing progress
        questions_found = []
        answers_found = []
        evidence_found = []

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Check for question number - handle multiple formats
            # Format 1: "Q1: Instructor Name"
            # Format 2: "1. **Question number (Q1)**: Q1"
            # Format 3: "**Question number (Q1)**: Q1"
            if 'Q' in line and ':' in line:
                # Try to extract Q number
                import re
                q_match = re.search(r'\b(Q\d+)\b', line)
                if q_match:
                    q_part = q_match.group(1)
                    if q_part in INVENTORY_QUESTIONS:
                        questions_found.append(q_part)

                        # If we have evidence but no answer for the previous question, try to infer answer
                        if current_question and current_question not in answers and current_question in evidence:
                            inferred_answer = infer_answer_from_evidence(current_question, evidence[current_question])
                            if inferred_answer:
                                answers[current_question] = inferred_answer

                        current_question = q_part
                        current_evidence = None

            # Check for evidence line - handle markdown bold
            # Format 1: "Evidence: text"
            # Format 2: "**Evidence**: text"
            elif 'Evidence' in line and ':' in line and current_question:
                # Remove markdown and extract evidence
                current_evidence = line.split(':', 1)[1].strip() if ':' in line else ''
                current_evidence = current_evidence.replace('**', '').strip()
                if current_evidence:
                    evidence[current_question] = current_evidence
                    evidence_found.append(current_question)

                    # If we have evidence but no answer yet, try to infer answer
                    if current_question not in answers:
                        inferred_answer = infer_answer_from_evidence(current_question, current_evidence)
                        if inferred_answer:
                            answers[current_question] = inferred_answer

            # Check for answer line - handle markdown bold
            # Format 1: "Answer: text"
            # Format 2: "**Answer**: text"
            elif 'Answer' in line and ':' in line and current_question:
                # Remove markdown and extract answer
                answer_text = line.split(':', 1)[1].strip() if ':' in line else ''
                answer_text = answer_text.replace('**', '').strip()
                if answer_text:
                    processed_answer = process_answer_text(current_question, answer_text)
                    if processed_answer:
                        answers[current_question] = processed_answer
                        answers_found.append(current_question)

        # Show parsing results and SAVE TO SESSION STATE
        st.session_state.debug_questions_found = len(questions_found)
        st.session_state.debug_answers_found = len(answers_found)
        st.session_state.debug_evidence_found = len(evidence_found)
        st.session_state.debug_questions_list = questions_found[:10]

        st.write(f"- Questions found in response: {len(questions_found)}")
        st.write(f"- Answers extracted: {len(answers_found)}")
        st.write(f"- Evidence found: {len(evidence_found)}")

        if questions_found:
            st.write(f"- First few questions found: {questions_found[:10]}")
        else:
            st.error("⚠️ No questions found in expected format (Q1:, Q2:, etc.)")
            st.write("**First 10 non-empty lines of response:**")
            non_empty = [l.strip() for l in lines if l.strip()][:10]
            st.session_state.debug_first_lines = non_empty
            for idx, l in enumerate(non_empty, 1):
                st.text(f"{idx}. {l[:100]}")

    except Exception as e:
        st.error(f"Error parsing LLM response: {str(e)}")
    
    # Final pass: check for questions with evidence but no answers
    for q_id in evidence:
        if q_id not in answers:
            inferred_answer = infer_answer_from_evidence(q_id, evidence[q_id])
            if inferred_answer:
                answers[q_id] = inferred_answer

    # DEBUG: Final summary
    st.write("---")
    st.write("✅ **FINAL EXTRACTION SUMMARY**")
    st.write(f"- **Total answers extracted:** {len(answers)}")
    st.write(f"- **Total evidence items:** {len(evidence)}")

    if answers:
        st.success(f"Successfully extracted {len(answers)} answers!")
        with st.expander("View extracted answers (first 10)"):
            for q_id, answer in list(answers.items())[:10]:
                st.write(f"**{q_id}:** {answer}")
    else:
        st.error("❌ NO ANSWERS WERE EXTRACTED!")
        st.write("**Possible reasons:**")
        st.write("1. LLM didn't follow the expected format (Q1:, Answer:, Evidence:)")
        st.write("2. LLM couldn't find information in the document")
        st.write("3. Response format changed - check the 'First 50 lines' above")

    # Store evidence in session state
    st.session_state.evidence = evidence
    return answers

# LLM Request Function
def make_llm_request(file_content1, filename1, file_content2=None, filename2=None):
    """Make LLM request with support for one or two documents"""
    url = "https://api.openai.com/v1/chat/completions"

    try:
        API_KEY = st.secrets["OPENAI_API_KEY"]
    except KeyError:
        st.error("OpenAI API key not found in secrets. Please configure your secrets.toml file.")
        return None

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }

    # Prepare document content
    documents_text = f"=== DOCUMENT 1 (Filename: {filename1}) ===\n" + file_content1
    if file_content2 and filename2:
        documents_text += f"\n\n=== DOCUMENT 2 (Filename: {filename2}) ===\n" + file_content2

    # Use the existing INVENTORY_PROMPT template
    prompt = INVENTORY_PROMPT.format(documents=documents_text)

    # DEBUG: Show prompt details and save to session state
    debug_info = []
    debug_info.append("🔍 **DEBUG: Request Details**")
    debug_info.append(f"- Document 1 length: {len(file_content1)} characters")
    if file_content2:
        debug_info.append(f"- Document 2 length: {len(file_content2)} characters")
    debug_info.append(f"- Total prompt length: {len(prompt)} characters")

    st.session_state.debug_request = "\n".join(debug_info)
    st.session_state.debug_prompt = prompt

    for line in debug_info:
        st.write(line)

    with st.expander("📄 View Full Prompt Being Sent to LLM"):
        st.text_area("Prompt", prompt, height=200)

    messages = [
        {
            "role": "user",
            "content": prompt
        }
    ]

    payload = {
        "model": "gpt-4o-mini",
        "messages": messages,
        "temperature": 0.3,
        "max_tokens": 4096
    }

    # DEBUG: Show what model we're using
    st.session_state.debug_model = payload["model"]

    try:
        with st.spinner('Analyzing document(s) and matching to inventory questions...'):
            response = requests.post(url, headers=headers, data=json.dumps(payload))

            st.write("🔍 **DEBUG: API Response**")
            st.write(f"- Status code: {response.status_code}")

            if response.status_code == 200:
                response_data = response.json()
                st.session_state.debug_status = response.status_code
                st.session_state.debug_response = str(response_data)[:500]

                # OpenAI standard response parsing
                content = response_data["choices"][0]["message"]["content"]

                st.write(f"- Response length: {len(content)} characters")
                st.write(f"- Model used: {response_data.get('model', 'unknown')}")
                st.write(f"- Tokens used: {response_data.get('usage', {})}")

                # SAVE TO SESSION STATE
                st.session_state.debug_llm_response = content
                st.session_state.debug_response_length = len(content)
                st.session_state.debug_model_used = response_data.get('model', 'unknown')
                st.session_state.debug_tokens = response_data.get('usage', {})

                # Show first 50 lines of response
                st.write("📝 **First 50 lines of LLM response:**")
                lines = content.split('\n')
                st.text_area("Response Preview", '\n'.join(lines[:50]), height=300)

                with st.expander("📄 View Full LLM Response"):
                    st.text_area("Full Response", content, height=400, key="full_response")

                return content
            else:
                st.error(f"Request failed with status code {response.status_code}")
                st.error(response.text)
                return None
    except requests.exceptions.RequestException as e:
        st.error(f"An error occurred: {str(e)}")
        return None
    except (KeyError, IndexError) as e:
        st.error(f"Error parsing OpenAI response: {str(e)}")
        st.error(f"Response data: {response_data}")
        return None

# UI Components
def create_input_widget(question_id, question_info, current_value=None):
    format_type = question_info["format"]

    if question_id in ["Q45", "Q46"]:
        options = ["not applicable", "no", "yes"]
        value = current_value if current_value in options else None

        index = options.index(value) if value in options else None
        return st.radio(
            question_info["question"],
            options=options,
            index=index,
            key=f"input_{question_id}",
            horizontal=True
        )

    if format_type == "y/n":
        return st.radio(
            question_info["question"],
            options=["No", "Yes"],
            index=0 if current_value == "No" else 1 if current_value == "Yes" else None,
            horizontal=True,
            key=f"input_{question_id}"
        )
    elif format_type.startswith("choice:"):
        options = format_type.split(":")[1].strip().split("/")
        index = options.index(current_value) if current_value in options else None
        return st.radio(
            question_info["question"],
            options=options,
            index=index,
            horizontal=True,
            key=f"input_{question_id}"
        )
    elif format_type == "percentage (0 to 100)":
        return st.number_input(
            question_info["question"],
            min_value=0,
            max_value=100,
            value=int(current_value) if current_value is not None and current_value != "" else None,
            key=f"input_{question_id}"
        )
    elif format_type == "number (minutes)" or format_type == "number":
        return st.number_input(
            question_info["question"],
            min_value=0,
            value=int(current_value) if current_value is not None and current_value != "" else None,
            key=f"input_{question_id}"
        )
    else:
        return st.text_input(
            question_info["question"],
            value=str(current_value) if current_value is not None else "",
            key=f"input_{question_id}"
        )

def is_valid_evidence(evidence_text):
    """Check if evidence is actually valid or just the LLM saying 'not found'"""
    if not evidence_text:
        return False

    # Phrases that indicate the LLM couldn't find information
    invalid_phrases = [
        "not applicable",
        "no mention",
        "not found",
        "no information",
        "could not find",
        "couldn't find",
        "no evidence",
        "not stated",
        "not specified",
        "does not mention",
        "doesn't mention"
    ]

    evidence_lower = evidence_text.lower()
    return not any(phrase in evidence_lower for phrase in invalid_phrases)

def display_section(section_name, question_ids, current_answers):
    """Display a section of questions with appropriate input widgets"""
    st.subheader(section_name)

    section_answers = {}
    all_answered = True

    # Create a container for this section
    section_container = st.container()

    with section_container:
        for q_id in question_ids:
            question_info = INVENTORY_QUESTIONS[q_id]
            current_value = current_answers.get(q_id)

            # Only use current_value if there's valid evidence for it
            if 'evidence' not in st.session_state:
                st.session_state.evidence = {}

            # Check if evidence exists AND is valid (not just "no mention found")
            has_valid_evidence = (q_id in st.session_state.evidence and
                                 is_valid_evidence(st.session_state.evidence.get(q_id)))

            # If there's a current value but no valid evidence, don't pre-fill
            if current_value and not has_valid_evidence:
                current_value = None

            # Use consistent column layout for all questions
            cols = st.columns([3, 1, 1])

            # Question and input widget
            with cols[0]:
                print(q_id)
                answer = create_input_widget(q_id, question_info, current_value)
                section_answers[q_id] = answer
                
                if answer is None or (isinstance(answer, str) and answer.strip() == ""):
                    all_answered = False
            
            # Status indicator
            with cols[1]:
                if current_value:
                    st.success("Pre-filled / check and change as needed")
                else:
                    if answer is None or (isinstance(answer, str) and answer.strip() == ""):
                        st.warning("Needs answer")
                    else:
                        st.success("Thank you")
            
            # Evidence expander
            with cols[2]:
                if 'evidence' not in st.session_state:
                    st.session_state.evidence = {}

                with st.expander("Why I selected this?"):
                    evidence_text = st.session_state.evidence.get(q_id, "")
                    if evidence_text and is_valid_evidence(evidence_text):
                        st.markdown("*Based on this text from your document:*")
                        st.write(evidence_text)
                    elif evidence_text and not is_valid_evidence(evidence_text):
                        st.write("LLM couldn't find relevant information in the document.")
                        st.caption(f"_LLM response: {evidence_text}_")
                    else:
                        st.write("No automated analysis available for this question.")
    
    return section_answers, all_answered


def process_sections(analyzed_answers):
    """Process each section of questions sequentially"""
    # Initialize session state variables
    if 'current_section' not in st.session_state:
        st.session_state.current_section = 0
    
    if 'all_answers' not in st.session_state:
        st.session_state.all_answers = analyzed_answers or {}
        
    if 'completed' not in st.session_state:
        st.session_state.completed = False
        
    if 'saved_to_sheets' not in st.session_state:
        st.session_state.saved_to_sheets = False
    
    # If already completed, show completion page
    if st.session_state.completed:
        st.title("Thank you for completing the inventory!")
        st.success("Your responses have been successfully uploaded to our database.")
        
        # Optional comments section
        # st.markdown("---")
        # comments = st.text_area(
        #     "Do you have any comments or want to mention other teaching practices that you use in this course?",
        #     height=150,
        #     key="optional_comments"
        # )
        
        # Update answers with comments if provided
        # if comments:
        #     st.session_state.all_answers["Q53"] = comments
        #     # Only save if we haven't saved before
        #     if not st.session_state.saved_to_sheets:
        #         save_to_google_sheets(st.session_state.all_answers)
        #         st.session_state.saved_to_sheets = True
        
        st.markdown("---")
        
        suggestions_wanted = st.radio(
            "Would you like ideas on what could be added to your next syllabus for this course?",
            options=["Yes", "No"],
            index=None,
            horizontal=True,
            key="suggestions_radio"
        )
        
        if suggestions_wanted == "Yes":
            suggestions = generate_syllabus_suggestions(st.session_state.all_answers)
            if suggestions:
                st.markdown("### Notes for future syllabus development, based on information you provided but was not in your syllabus")
                st.markdown(suggestions)
                st.info("You can copy these notes for future use when developing your syllabus.")
        # elif suggestions_wanted == "No":
        st.success("Thank you for participating! You may now close this browser - you results are already saved.")
        
        return
    
    # Regular section processing for incomplete form
    main_container = st.container()
    
    with main_container:
        sections_list = list(SECTIONS.items())
        current_section = sections_list[st.session_state.current_section]
        
        section_name, question_ids = current_section
        section_answers, all_answered = display_section(section_name, question_ids, st.session_state.all_answers)
        
        # Navigation buttons
        nav_cols = st.columns(3)
        
        with nav_cols[0]:
            if st.session_state.current_section > 0:
                if st.button("Previous Section"):
                    st.session_state.all_answers.update(section_answers)
                    st.session_state.current_section -= 1
                    st.rerun()
        
        with nav_cols[1]:
            st.write(f"Section {st.session_state.current_section + 1} of {len(SECTIONS)}")
        
        with nav_cols[2]:
            if st.session_state.current_section < len(SECTIONS) - 1:
                if all_answered:
                    if st.button("Next Section"):
                        st.session_state.all_answers.update(section_answers)
                        st.session_state.current_section += 1
                        st.rerun()
                else:
                    st.warning("Please answer all questions")
            elif all_answered:
                # Update answers before showing complete button
                st.session_state.all_answers.update(section_answers)
                
                # Complete button and subsequent actions
                if st.button("Complete"):
                    # Only save if we haven't saved before
                    if save_to_google_sheets(st.session_state.all_answers):
                        st.session_state.completed = True
                        st.session_state.saved_to_sheets = True
                        st.rerun()
                    else:
                        st.error("Could not save to Google Sheets")
                else:
                    st.warning("Please click Complete to finish and save your responses")


def process_answer_text(question_id, answer_text):
    """Process answer text based on question format"""
    if not answer_text:
        return None
        
    q_format = INVENTORY_QUESTIONS[question_id]["format"]
    answer_text = answer_text.lower()
    
    try:
        if q_format == "y/n":
            # Extended yes/no detection
            yes_indicators = ['yes', 'does', 'do', 'provides', 'has', 'used', 'given', 'shown', 
                            'true', 'correct', 'will', 'available', 'included', 'required']
            no_indicators = ['no', 'does not', "doesn't", 'do not', "don't", 'not used', 'not given', 
                           'false', 'unavailable', 'excluded', 'none']
            
            if any(indicator in answer_text for indicator in yes_indicators):
                return "Yes"
            elif any(indicator in answer_text for indicator in no_indicators):
                return "No"
                
        elif q_format.startswith("choice:"):
            options = q_format.split(":")[1].strip().split("/")
            
            # Special handling for frequency-based questions
            if question_id in ["Q23", "Q25", "Q26", "Q27"]:
                if "never" in answer_text or "don't" in answer_text or "not used" in answer_text:
                    return "never"
                elif "every class" in answer_text or "each class" in answer_text:
                    return "every class"
                elif "week" in answer_text:
                    return "every week"
                elif "rarely" in answer_text or "seldom" in answer_text:
                    return "rarely"
                elif "sometimes" in answer_text or "occasionally" in answer_text:
                    return "once in a while"
            
            # Handle Q43 (trying new methods)
            elif question_id == "Q43":
                if "never" in answer_text or "don't" in answer_text:
                    return "never"
                elif "often" in answer_text or "frequently" in answer_text:
                    return "often"
                elif "rarely" in answer_text or "seldom" in answer_text:
                    return "rarely"
                elif "sometimes" in answer_text or "occasionally" in answer_text:
                    return "sometimes"
            
            # Handle TA questions (Q45, Q46)
            elif question_id in ["Q45", "Q46"]:
                if "not applicable" in answer_text or "n/a" in answer_text:
                    return "not applicable"
                elif any(yes_word in answer_text for yes_word in ['yes', 'do', 'does', 'provided']):
                    return "yes"
                elif any(no_word in answer_text for no_word in ['no', 'don\'t', 'does not']):
                    return "no"
            
            # Default option matching
            for option in options:
                if option.lower() in answer_text:
                    return option
                    
        elif q_format in ["percentage (0 to 100)", "number (minutes)", "number"]:
            import re
            numbers = re.findall(r'\d+', answer_text)
            if numbers:
                num = int(numbers[0])
                if q_format == "percentage (0 to 100)":
                    return min(100, max(0, num))  # Clamp between 0 and 100
                else:
                    return num
                    
        else:  # text format
            return answer_text.strip()
            
    except Exception as e:
        st.error(f"Error processing answer text for {question_id}: {str(e)}")
    
    return None

def infer_answer_from_evidence(question_id, evidence_text):
    """Attempt to infer an answer from evidence text"""
    if not evidence_text:
        return None
        
    try:
        q_format = INVENTORY_QUESTIONS[question_id]["format"]
        evidence_text = evidence_text.lower()
        
        if q_format == "y/n":
            # Look for positive indicators in evidence
            positive_indicators = [
                'yes', 'does', 'do', 'provides', 'has', 'used', 'given', 'shown',
                'available', 'included', 'required', 'expected', 'will be', 'must',
                'students should', 'students will', 'students are'
            ]
            # Look for negative indicators in evidence
            negative_indicators = [
                'no', 'not', "don't", 'does not', "doesn't", 'do not',
                'unavailable', 'excluded', 'forbidden', 'prohibited'
            ]
            
            # Check for positive indicators first
            if any(indicator in evidence_text for indicator in positive_indicators):
                return "Yes"
            # Then check for negative indicators
            elif any(indicator in evidence_text for indicator in negative_indicators):
                return "No"
                
        elif q_format.startswith("choice:"):
            options = q_format.split(":")[1].strip().split("/")
            # Look for option matches in evidence
            for option in options:
                if option.lower() in evidence_text:
                    return option
                    
        elif q_format in ["percentage (0 to 100)", "number (minutes)", "number"]:
            import re
            numbers = re.findall(r'\d+', evidence_text)
            if numbers:
                num = int(numbers[0])
                if q_format == "percentage (0 to 100)" and num <= 100:
                    return num
                elif q_format == "number (minutes)" or q_format == "number":
                    return num
                    
        else:  # text format
            # For text format, extract meaningful content from evidence
            # Remove common quote markers and clean up
            cleaned_text = evidence_text.replace('"', '').replace('"', '').strip()
            if len(cleaned_text) > 0:
                return cleaned_text
                
    except Exception as e:
        st.error(f"Error inferring answer from evidence for {question_id}: {str(e)}")
    
    return None

def reset_form():
    """Reset all form state"""
    for key in ['analyzed_answers', 'current_section', 'all_answers', 'evidence']:
        if key in st.session_state:
            del st.session_state[key]

def make_llm_request_for_suggestions(prompt):
    """Make LLM request specifically for syllabus suggestions"""
    url = "https://api.openai.com/v1/chat/completions"

    try:
        API_KEY = st.secrets["OPENAI_API_KEY"]
    except KeyError:
        st.error("OpenAI API key not found in secrets.")
        return None

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }

    messages = [
        {
            "role": "user",
            "content": prompt
        }
    ]

    payload = {
        "model": "gpt-4o-mini",
        "messages": messages,
        "temperature": 0.7,  # Slightly higher temperature for more creative suggestions
        "max_tokens": 4096
    }

    try:
        with st.spinner('Generating syllabus suggestions...'):
            response = requests.post(url, headers=headers, data=json.dumps(payload))

            if response.status_code == 200:
                response_data = response.json()
                return response_data["choices"][0]["message"]["content"]
            else:
                st.error(f"Request failed with status code {response.status_code}")
                return None
    except requests.exceptions.RequestException as e:
        st.error(f"An error occurred: {str(e)}")
        return None
    except (KeyError, IndexError) as e:
        st.error(f"Error parsing OpenAI response: {str(e)}")
        return None

def generate_syllabus_suggestions(answers):
    """Generate suggestions for syllabus improvements based on inventory answers"""
    
    # Questions to skip (not typically in syllabus or practice-based)
    SKIP_QUESTIONS = set([f"Q{i}" for i in range(1, 5)] +  # Basic info
                        [f"Q{i}" for i in range(19, 25)] +  # In-class practices
                        ["Q33"] +  # Feedback practices
                        [f"Q{i}" for i in range(44, 52)])  # TA and collaboration info
    
    # Create prompt for LLM
    prompt = """You are an expert in the best practices for syllabus development in higher education.
    Based on these teaching inventory answers, provide sample syllabus text that is clear, 
    direct, and student-friendly. For each section:
    1. Provide ready-to-use syllabus language in plain English (avoid jargon and words like "utilize")
    2. Add brief notes about what specific details could be added that you do not have
    
    Organize the suggestions into these syllabus sections:
    - Course Materials and Resources
    - Learning Objectives and Outcomes
    - Course Policies (including AI/LLM policies)
    - Assignments and Assessments
    - Student Support and Resources
    - Course Communication
    
    Format your response as:
    [Section Name]
    
    [Sample syllabus text]
    
    Notes: [What specific information should be added]
    
    ---
    
    Current Teaching Practices (based on inventory answers):
    """
    
    # Add relevant answers to prompt
    suggestions_needed = []
    for q_id, answer in answers.items():
        if q_id not in SKIP_QUESTIONS and answer:  # Only include questions with answers
            question_text = INVENTORY_QUESTIONS[q_id]["question"]
            if answer == "Yes":
                suggestions_needed.append(f"✓ {question_text}")
            elif answer == "No":
                suggestions_needed.append(f"○ {question_text}")
            else:
                suggestions_needed.append(f"• {question_text}: {answer}")
    
    prompt += "\n".join(suggestions_needed)
    
    # Add specific guidance for the response
    prompt += """

    For your response:
    1. Write in clear, simple language that students will easily understand
    2. Avoid academic jargon and complex terminology
    3. Focus on what will be done in the course
    4. Include notes about specific details the instructor should add
    5. Present the text exactly as it could appear in a syllabus
    6. Don't suggest anything about "Contact information" or "Student Support and Resources"
    
    Do not include explanations - only provide the sample syllabus text and notes for customization."""

    return make_llm_request_for_suggestions(prompt)
    


def main():
    st.set_page_config(layout="wide")

    # VERSION INFO - Always visible
    st.sidebar.write("# 📌 VERSION INFO")
    st.sidebar.success("**Version 2.5** - SECTIONS REBALANCED")
    st.sidebar.write("✓ Q5, Q10, Q44 - Not Asked (show 'notasked')")
    st.sidebar.write("✓ Q38 - Changed to frequency choice")
    st.sidebar.write("✓ Only pre-fill if VALID evidence found")
    st.sidebar.write("✓ 8 balanced sections (6-7 questions each)")
    st.sidebar.write("✓ Questions kept in original order")
    st.sidebar.write("---")

    # Show persistent debug info in sidebar
    if 'debug_model' in st.session_state:
        st.sidebar.write("# 🔍 DEBUG INFO")

        with st.sidebar.expander("📊 Extraction Results", expanded=True):
            st.write(f"**Questions found:** {st.session_state.get('debug_questions_found', 'N/A')}")
            st.write(f"**Answers extracted:** {st.session_state.get('debug_answers_found', 'N/A')}")
            st.write(f"**Evidence items:** {st.session_state.get('debug_evidence_found', 'N/A')}")

            if 'debug_questions_list' in st.session_state and st.session_state.debug_questions_list:
                st.write(f"**Questions:** {', '.join(st.session_state.debug_questions_list)}")

        with st.sidebar.expander("🤖 API Details"):
            st.write("**Model:**", st.session_state.debug_model)
            st.write("**Status:**", st.session_state.get('debug_status', 'N/A'))
            if 'debug_response_length' in st.session_state:
                st.write("**Response length:**", st.session_state.debug_response_length)
            if 'debug_tokens' in st.session_state:
                st.write("**Tokens:**", st.session_state.debug_tokens)

        with st.sidebar.expander("📝 LLM Response Preview"):
            if 'debug_llm_response' in st.session_state:
                lines = st.session_state.debug_llm_response.split('\n')[:20]
                st.text('\n'.join(lines))
            else:
                st.write("No response captured")

        with st.sidebar.expander("⚠️ First Lines (if parsing failed)"):
            if 'debug_first_lines' in st.session_state:
                for idx, line in enumerate(st.session_state.debug_first_lines, 1):
                    st.text(f"{idx}. {line[:80]}")
    
    # Initialize session state
    if 'analyzed_answers' not in st.session_state:
        st.session_state.analyzed_answers = None
    
    if 'evidence' not in st.session_state:
        st.session_state.evidence = {}
        
    if 'started' not in st.session_state:
        st.session_state.started = False

    # Only show upload page if not started
    if not st.session_state.started:
        st.title("Vanderbilt Psychology Teaching Inventory Helper")


        
        st.write("""
        This is an inventory of teaching practices as they apply for a specific course in a 
        specific semester. This information is collected only to help us better understand how we teach 
        our courses (it will not be used for any evaluation). The Full inventory has several
        questions but this tool will attempt to help you answer them as fast as possible.
        
        The inventory lists several evidence-based practices that are unlikely to all be used in the same
        course. We know practices differ based on class size and type of content. Please answer conservatively, 
        as we are looking for an accurate picture of what practices are really used, not an ideal picture.

        If you have a syllabus for the course, or any other document relevant to your teaching practices 
        in this course, please upload it (in pdf or .docx format). If you don't, you can answer all questions manually. 
        At the end, we will provide you with suggestions for information to add to your syllabus.
                 
        This app uses Generative AI              
        """)

        
        # File uploaders in columns
        col1, col2 = st.columns(2)
        with col1:
            file1 = st.file_uploader("Upload your syllabus", 
                                    type=["pdf", "docx"],
                                    key="file1")
        with col2:
            file2 = st.file_uploader("Upload any additional teaching document (optional)", 
                                    type=["pdf", "docx"],
                                    key="file2")
        
        # Start button
        if st.button("Start with these documents"):
            with st.spinner('Wait while AI extracts information...'):
                # Process uploaded files
                content1, filename1 = process_uploaded_file(file1) if file1 else (None, None)
                content2, filename2 = process_uploaded_file(file2) if file2 else (None, None)
                
                if content1 or file1 is None:
                    if content1:
                        st.success("Document(s) processed successfully!")
                    
                    if st.session_state.analyzed_answers is None:
                        if content1:
                            response = make_llm_request(content1, filename1, content2, filename2)
                            if response:
                                st.session_state.analyzed_answers = parse_llm_response(response)
                        else:
                            st.session_state.analyzed_answers = {}
                    
                    # Set started to True and rerun to clear the page
                    st.session_state.started = True
                    st.rerun()
    
    # Show sections if we're started or have answers
    else:
        if st.session_state.analyzed_answers is not None or 'current_section' in st.session_state:
            process_sections(st.session_state.analyzed_answers)

if __name__ == "__main__":
    main()
